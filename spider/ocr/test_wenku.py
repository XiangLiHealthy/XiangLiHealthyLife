#usr/bin/env python
#coding=utf-8
from docx.shared import RGBColor
from requests import get
from requests.exceptions import ReadTimeout
from chardet import detect
from bs4 import BeautifulSoup, Tag
from os import getcwd,mkdir
from os.path import join,exists
from re import findall
from json import loads
from time import time
import re

import os
import docx

class ParserLocalHtml:
    def __init__(self, dir,save) :
        self.dir = dir
        self.save_dir = save

    def getHtmlList(self):
        #扫描该目录下的所有html文件
        files_list = []
        postfix = '.html'
        prefix = 'xx'

        for root, sub_dirs, files in os.walk(self.dir):
            for special_file in files:
                if postfix:
                    if special_file.endswith(postfix):
                        files_list.append(os.path.join(root, special_file))
                elif prefix:

                    if special_file.startswith(prefix):

                        files_list.append(os.path.join(root, special_file))
                    else:
                        files_list.append(os.path.join(root, special_file))

        return files_list

    def getHtmlContent(self, path):
        with open(path, 'rb+') as f:
            content = f.read()
            return content

        return None

    def parserHtml(self, content):
        result = {}

        #soup
        soup = BeautifulSoup(content, 'html.parser')

        #get title content
        result['title'] = soup.findAll('h3', attrs={'class': 'doc-title'})[0]
        #result['title'] = title[0].contents[0]

        lines = soup.findAll('p', attrs = {'class' : re.compile('reader-word-layer|reader-pic-item')})
        result['body'] = lines

        return result
    def save(self, content):
        try:
            file = self.save_dir + '/' + content['title'].get_text() + '.docx'
            doc = docx.Document()

            #写入标题
            doc.add_heading(content['title'].get_text())

            #写入内容
            paragraph = doc.add_paragraph()
            r = paragraph.add_run()
            for tag in content['body'] :
                if 'reader-word-layer' == tag.attrs['class'][0] :
                    r = self.save_txt(paragraph, tag)

                elif 'reader-pic-item' == tag.attrs['class'][0] :
                    self.save_pic(doc, paragraph, tag)
                else:
                    print('unknow tag class:{}'.format(tag.attrs['class'][0]))

            doc.save(file)

        except Exception as e:
            print (e)


    def save_txt(self, doc, tag):
        txt = tag.get_text()
        attrs = tag.attrs

        r = doc.add_run(txt)
        return r

    def save_pic(self, doc, paragraph, tag):
        try:
            for item in tag.contents:
                if isinstance(item, Tag) :
                    pic = item.attrs
                    doc.add_picture('/home/xiangbaosong/jianhai/spider/html/' + pic['src'])
                else:
                    paragraph.add_run('***').font.color.rgb = RGBColor(250,0,0)
        except Exception as e :
            print(e)
        return

    def performSpider(self):
        htmlList = self.getHtmlList()

        for html in htmlList :
            try:

                src = self.getHtmlContent(html)
                txt = self.parserHtml(src)
                self.save(txt)

            except Exception as e :
                print ('{},html:{}', e, html)

            print (txt)

        return


if __name__ == '__main__':
    spider = ParserLocalHtml('/home/xiangbaosong/jianhai/spider/html', '/home/xiangbaosong/jianhai/spider')
    spider.performSpider()

