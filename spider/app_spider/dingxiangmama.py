#

import csv
import io
import json

import requests
import chardet
import urllib.parse
from bs4 import BeautifulSoup
import os
import downloader
import cv2
from io import BytesIO
from PIL import Image
import docx
#from docx.enum.style import WD_STYLE_TYP, WD_STYLE_TYPE

g_header = {}
g_header['app-ac'] = '3e1876d3-d315-4e15-8350-e9a2ed0576ac'
g_header['app-mc'] = '357e50ed621645e9908be1e308d9ecf5'
g_header['app-version'] = '4.17.0'
g_header['app-os_version'] = '6.0.1'
g_header['app-ssion-id'] = 'ce422064-669a-4d11-abd0-d4a4d90e7983'
g_header['user-agent'] = 'Mozilla/5.0 (Linux; Android 6.0.1; DUK-AL20 Build/V417IR; wv) AppleWebKit/537.36 (KHTML, like Gecko) Version/4.0 Chrome/52.0.2743.100 Mobile Safari/537.36 dxyapp_name/gaia dxyapp_version/4.17.0 dxyapp_system_version/6.0.1 dxyapp_client_id/357e50ed621645e9908be1e308d9ecf5 dxyapp_ac/3e1876d3-d315-4e15-8350-e9a2ed0576ac dxyapp_sid/b16803e2-7044-4df4-a5c0-5ae19bb7a3ba'
g_header['accept-encoding'] = 'gzip'
g_header['cookie'] = 'DXY_CHD_SESSION=eyJhIjoxMjIxNzgyNDEyLCJ0IjoxNjEyNDkwMDAzLCJuIjoieEo1M1djNE9xeHo0eGJrSSIsImQiOiJ7XCJhdHRyaWJ1dGVzXCI6e1wic3NvXCI6XCJkeHlfcGh0NWR4eDVcIixcInZcIjo1OSxcIm1JZFwiOjM0NTc2Njg1Njc0MzYzMjE0Nzh9LFwiaWRcIjozNDU3NjY4MTMxNDk3MTM4MTg2LFwidXNlcm5hbWVcIjpcImR4eV9waHQ1ZHh4NVwiLFwibWFya3NcIjo1MTQsXCJtb21cIjoxMjAyMTAyMDh9IiwicyI6IjE3MGEyMmIxMDBiZWE1ZmI4ZWIwMzdhYTk4NmJhZDIyZTZhMTQxZTMifQ==; route=0fd188492be2c1bd7c318c8297006bb8; DXY_TRACE_ID=ZnuBRXuKLRdnprhKLjHcYrqg9Ovts0Ai'


g_path = 'D:/jianhai/spider/data'

class Calender :
    def __init__(self):
        # 1.孕期日历
        self.baby_change_api = 'https://mama.dxy.com/japi/platform/201520012?weeks='
        self.mather_change_api = 'https://mama.dxy.com/japi/platform/201221150?topEntityId=3370220475686096408&topEntityType=6&checkRead=false'
        self.father_advise_api = 'https://mama.dxy.com/japi/platform/201221150?topEntityId=3370220948132499836&topEntityType=6&checkRead=false'

        self.downloader = downloader.Downloader()
        self.path = 'D:/jianhai/spider/data'

    def get_calender(self, weeks):
        #1.construct header
        url = '{}{}'.format(self.baby_change_api, weeks)

        #2.perform http
        res = self.downloader.get(url, g_header)

        #3.parse parmeter
        js = json.loads(res)
        items = js['results']['item']
        return items
    def save_img(self, path, content):
        try:
            with open(path, "wb") as f:
                f.write(content)

        except Exception as e :
            print (e)

        return

    def get_imgs(self, weeks, color_drapper_url, fruit_url):

        img_header = {}
        img_header['if-modified-since'] = 'Mon, 19 Aug 2019 02:51:42 GMT'
        img_header['user-agent'] = 'Mozilla/5.0 (Linux; Android 6.0.1; DUK-AL20 Build/V417IR; wv) AppleWebKit/537.36 (KHTML, like Gecko) Version/4.0 Chrome/52.0.2743.100 Mobile Safari/537.36 dxyapp_name/gaia dxyapp_version/4.17.0 dxyapp_system_version/6.0.1 dxyapp_client_id/357e50ed621645e9908be1e308d9ecf5 dxyapp_ac/3e1876d3-d315-4e15-8350-e9a2ed0576ac dxyapp_sid/b16803e2-7044-4df4-a5c0-5ae19bb7a3ba'
        img_header['if-none-match'] = '46B2C45E195A2D6B2352F64084574834'
        img_header['accept'] = 'image/webp,image/*,*/*;q=0.8'
        img_header['accept-encoding'] = 'gzip, deflate'
        img_header['accept-language'] = 'zh-CN,en-US;q=0.8'
        img_header['x-requested-with'] = 'com.dxy.gaia'

        images = {}
        color = self.downloader.get(color_drapper_url, img_header, True)
        fruit = self.downloader.get(fruit_url, img_header, True)

        color_drapper_file = '{}/孕期日历婴儿图/baby_{}.png'.format(self.path, weeks)
        fruit_file = '{}/孕期日历婴儿图/fruit_{}.png'.format(self.path, weeks)
        images['color_drapper'] = ('baby_{}.png'.format(weeks))
        images['fruit'] = ('fruit_{}.png'.format( weeks))

        self.save_img(color_drapper_file, color)
        self.save_img(fruit_file, fruit)


        return images

    def get_article(self, id):
        url = 'https://mama.dxy.com/japi/platform/201221150?topEntityId={}&topEntityType=6&checkRead=false'.format(id)
        res = self.downloader.get(url, g_header)
        js = json.loads(res)

        items = js['results']['items'][0]
        contents = []
        for item in items['childModules'] :
            richText = item['richText']
            contents.append(richText['cleanContent'])

        return contents

    def write_to_docx(self, weeks, imgs, mom_article, dad_article):
        #create file if not exist and write header
        file = '{}/孕期日历.docx'.format(self.path)
        if not os.path.exists(file) :
            doc = docx.Document()
            table = doc.add_table(1, 9, style = 'Table Grid')
            row = table.rows[0].cells
            row[0].text = '第几周'
            row[1].text = '婴儿重量g'
            row[2].text = '婴儿身高'
            row[3].text = '描述'
            row[4].text = '每天变化'
            row[5].text = '妈妈变化'
            row[6].text = '爸爸建议'
            row[7].text = '婴儿图片'
            row[8].text = '水果类比'


            doc.save(file)

        #write content
        doc = docx.Document(file)
        table = doc.tables[0]
        row = table.add_row().cells

        try:
            row[0].text = str(weeks['weeks'])
            row[1].text = weeks['weight']
            row[2].text = weeks['height']
            row[3].text = weeks['description']

            #weeks 0,1,2,3,4,5,6
            for day in weeks['weekdays'] :
                new_row = row[4].add_table(1,1).rows[0].cells
                new_row[0].text = '周，距离出生日期{}天，描述：{}'.format(day['days'], day['countDownDay'], day['description'])

            #mon article
            for change in mom_article :
                new_row = row[5].add_table(1,1).rows[0].cells
                new_row[0].text = change

            #father article
            for advise in dad_article :
                new_row  = row[6].add_table(1,1).rows[0].cells
                new_row[0].text = advise

            row[7].paragraphs[0].add_run().add_picture(imgs['color_drapper'], width=50)
            row[8].paragraphs[0].add_run().add_picture(imgs['fruit'], width=50)

            doc.save(file)
        except Exception as e :
            print (e)
        return

    def save_csv(self, weeks, imgs, mom_article, dad_article):
        # create file name
        file_name = '{}/孕期日历.csv'.format(self.path)
        mode = ''
        if not os.path.exists(file_name):
            header = ['第几周', '婴儿重量g', '婴儿身高', '描述', '每天变化', '妈妈变化', '爸爸建议', '婴儿图片', '类比水果']
            with open(file_name, 'w', encoding='utf-8-sig', newline='') as f:
                writer = csv.writer(f)
                writer.writerow(header)

        with open(file_name, 'a+', encoding='utf-8-sig', newline='') as f:
            try:
                data = []
                data.append(str(weeks['weeks']))
                data.append(weeks['weight'])
                data.append(weeks['height'])
                data.append(weeks['description'])

                tmp = ''
                # weeks 0,1,2,3,4,5,6
                for day in weeks['weekdays']:
                    tmp += '周{}，距离出生日期{}天，描述：{}'.format(day['days'] + 1, day['countDownDay'], day['description'])
                    tmp += '\n\n'
                data.append(tmp)

                # mon article
                tmp = ''
                for change in mom_article:
                    tmp += change
                    tmp += '\n\n'
                data.append(tmp)

                # father article
                tmp = ''
                for advise in dad_article:
                    tmp += advise
                    tmp += '\n\n'
                data.append(tmp)

                data.append(imgs['color_drapper'])
                data.append(imgs['fruit'])
                datas = [data]

                writer = csv.writer(f)
                writer.writerows(datas)
            except Exception as e:
                print(e)


        return

    def test_get_pic(self):
        url = 'https://img1.dxycdn.com/2019/1012/265/3373427606799737424-2.png'
        re = requests.get(url)
        with open('./test.png', 'wb')as f:
            f.write(re.content)

        return

    def perform_spider(self):
        self.test_get_pic()

        for idx in range(1,41) :
            try:
                #1.get weeks summary
                weeks = self.get_calender(idx)

                #2.get baby images
                imgs = self.get_imgs(idx, weeks['colorDopplerImg'], weeks['fruitsImg'])

                #3.get mother article
                mather_article = self.get_article(weeks['momArticleId'])

                #4.get dad article
                father_article = self.get_article(weeks['dadArticleId'])

                #5.write to docx
                self.save_csv(weeks, imgs, mather_article, father_article)

            except Exception as e:
                print(e)
        print ('calender speder completed.......')
        return

class ExaminationAlarm :
    def __init__(self):
        # 2.产检提醒
        self.downloader = downloader.Downloader()

    def get_items(self, url):
        # perform http
        res = self.downloader.get(url, g_header)

        # translate into json and get items
        js = json.loads(res)
        items = js['items']

        return items

    def get_examines(self):
        #create url
        url = 'https://mama.dxy.com/japi/platform/201520022?momId=3457668567436321478&dateOfDelivery=2021-02-08'

        # perform http
        res = self.downloader.get(url, g_header)

        # translate into json and get items
        js = json.loads(res)
        items = js['results']['items']

        return items

    def get_detail(self, id):
        #create url
        url = 'https://mama.dxy.com/japi/platform/201221150?topEntityId={}&topEntityType=6'.format(id)

        # perform http
        res = self.downloader.get(url, g_header)

        # translate into json and get items
        js = json.loads(res)
        items = js['results']['items']
        return items

    def get_description(self, id):
        #create url
        url = 'https://mama.dxy.com/japi/platform/201220081?id={}'.format(id)
        res = self.downloader.get(url, g_header)

        #get value
        js = json.loads(res)
        description = js['results']['item']['brief']

        return description

    def save_csv(self, examine, details):
        #if not exist create and write title
        # create file name
        file_name = '{}/丁香妈妈--产检提醒.csv'.format(g_path)
        mode = ''
        if not os.path.exists(file_name):
            header = ['第几次', '第几周', '产检概述', '产检准备', '重点项目', '常规项目', '是否空腹']
            with open(file_name, 'w', encoding='utf-8-sig', newline='') as f:
                writer = csv.writer(f)
                writer.writerow(header)

        with open(file_name, 'a+', encoding='utf-8-sig', newline='') as f:
            try:
                data = []
                data.append(examine['times'])
                data.append(examine['timeSlot'])

                for detail in details :
                    content = ''
                    title = detail['rootModule']['title']
                    content += title + '\n'

                    if '产检概述' == title :
                        content += self.get_description(detail['rootModule']['parentEntityId'])
                    else:
                        child = ''
                        for child_module in detail['childModules'] :
                            child_title = child_module['title']
                            child_text = child_module['richText']['cleanContent']
                            line = '{}\n{}'.format(child_title, child_text)
                            content += line + '\n\n'

                    content += '\n\n'
                    data.append(content)

                eat_flag = {}
                eat_flag['1'] = '是'
                eat_flag['0'] = '否'
                key = str(examine['eatFlag'])
                data.append(eat_flag[key])
                datas = [data]

                writer = csv.writer(f)
                writer.writerows(datas)
            except Exception as e:
                print(e)

        return

    def performSpider(self):
        #get examine list
        try:
            examines = self.get_examines()

            # get detail description by id
            for examine in examines :
                try:
                    detail = self.get_detail(examine['contentId'])

                    self.save_csv(examine, detail)#save to csv
                except Exception as e :
                    print (e)

            print ('completed spider counter:{}'.format(len(examines)))
        except Exception as e:
            print (e)

        return

class IndexExplain :
    def __init__(self):
        self.downloader = downloader.Downloader()

    def get_index_list(self):
        #create url
        url = 'https://mama.dxy.com/japi/platform/201221151?pgcCategoryId=3363761035998934193'

        #perform http
        res = self.downloader.get(url, g_header)

        #js value
        js = json.loads(res)
        indexs = js['results']['items']

        return indexs

    def get_detail(self, id):
        #create url
        url = 'https://mama.dxy.com/japi/platform/201221150?topEntityId={}&topEntityType=6'.format(id)

        #perform http
        res = self.downloader.get(url, g_header)

        #js value
        js = json.loads(res)
        content = js['results']['items'][0]['childModules'][0]
        html = content['richText']['content']
        soup = BeautifulSoup(html, 'html.parser')
        divs = soup.findAll('div')

        format_text = ''
        for div in divs :
            format_text += div.get_text()

        return format_text

    def save_csv(self, index, article, detail):
        # if not exist create and write title
        # create file name
        file_name = '{}/丁香妈妈--产检解读.csv'.format(g_path)
        mode = ''
        if not os.path.exists(file_name):
            header = ['项目', '分类', '产检概述', '详细描述']
            with open(file_name, 'w', encoding='utf-8-sig', newline='') as f:
                writer = csv.writer(f)
                writer.writerow(header)

        with open(file_name, 'a+', encoding='utf-8-sig', newline='') as f:
            try:
                data = []
                data.append(article['title'])
                data.append(index['categoryName'])
                data.append(article['brief'])
                data.append(detail)

                datas = [data]

                writer = csv.writer(f)
                writer.writerows(datas)
            except Exception as e:
                print(e)

    def performSpider(self):
        try:
            #get index list
            indexs = self.get_index_list()

            #recursive get detail
            for index in indexs :
                cmsArticles = index['cmsArticles']
                if len(cmsArticles) < 1 :
                    continue

                for article in cmsArticles :
                    detail = self.get_detail(article['id'])
                    self.save_csv(index, article, detail)

        except Exception as e :
            print (e)

        print ('spider index explain completed')

        return

class DXmama :
    def __init__(self):
        self.log_api = 'https://auth.dxy.cn/api/login/v2/phone'

        #3.产检解读
        self.unsramble_api = 'https://mama.dxy.com/japi/platform/201221151?pgcCategoryId=3363761035998934193'

    def get_classification(self):
        #perform http
        res = self.perform_http(self.classification_api)


        classification_list = []
        soup = json.loads(res)

        return soup['data']
    def get_category_list(self, category_id):
        url = '{}{}'.format(self.category_api, category_id)
        res = self.perform_http(url)
        js = json.loads(res)
        return js['data']

    def get_detail_list(self, action_id):
        detail_list = []
        url = '{}{}'.format(self.detail_api, action_id['id'])

        res = self.perform_http(url)
        soup = json.loads(res)

        return soup['data']
    def write_content(self, detail, type):
        #create file name
        file_name = 'D:\jianhai\spider\data\杭州健康通-能不能做.csv'
        mode = ''
        if not os.path.exists(file_name):
            header = ['name', 'title', 'type', 'description', 'canDo']
            with open(file_name, 'w', encoding='utf-8-sig', newline='') as f:
                writer = csv.writer(f)
                writer.writerow(header)

        with open(file_name, 'a+', encoding='utf-8-sig', newline='') as f:
            try:
                data = []
                data.append(detail['name'])
                data.append(detail['title'])
                data.append(type)
                data.append(detail['description'])

                canDo = '{}'.format(detail['canDo'])
                data.append(self.canDo[canDo])
                datas = [data]

                writer = csv.writer(f)
                writer.writerows(datas)
            except Exception as e:
                print (e)

        return

    def performSpider(self):
        #calender = Calender()
        #spider calender change
        #calender.perform_spider()
        #print ('perform calender completed')

        #spidder alarm
        #examine = ExaminationAlarm()
        #examine.performSpider()

        #spider unscrable
        indexExplain = IndexExplain()
        indexExplain.performSpider()

        return


if __name__ == '__main__' :
    spider = DXmama()
    spider.performSpider()

    print ('complete spider')